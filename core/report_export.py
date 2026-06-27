"""
检测报告导出 —— Word (.docx) 与 PDF。
输入为 /api/inspect 返回的完整 payload：{project_info, report, meta}，不重跑 LLM。
"""
import io
import os
from datetime import datetime

# ── 公共：严重程度展示 ─────────────────────────────────
_SEV = {
    "red":    ("否决项·必改", "F53F3F"),
    "yellow": ("扣分项·建议改", "FF7D00"),
    "blue":   ("待核验·人工", "722ED1"),
    "green":  ("合规", "00B42A"),
}


def _sev(sev: str):
    return _SEV.get(sev, ("提示", "86909C"))


def _overview_lines(payload: dict):
    pi = payload.get("project_info") or {}
    report = payload.get("report") or {}
    meta = payload.get("meta") or {}
    counts = report.get("counts") or {}
    return pi, report, meta, counts


# ════════════════════════════════════════════════════════
#  Word (.docx)
# ════════════════════════════════════════════════════════
def build_docx(payload: dict) -> bytes:
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn

    pi, report, meta, counts = _overview_lines(payload)

    def font(run, name="仿宋", size=12, bold=False, color=None):
        run.font.name = name
        run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
        run.font.size = Pt(size)
        run.font.bold = bold
        if color:
            run.font.color.rgb = RGBColor.from_string(color)

    doc = Document()
    sec = doc.sections[0]
    sec.page_height, sec.page_width = Cm(29.7), Cm(21)
    sec.left_margin = sec.right_margin = Cm(2.5)
    sec.top_margin = sec.bottom_margin = Cm(2.2)

    # 标题
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    font(p.add_run("标书规范性检测报告"), "黑体", 22, True)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    font(p.add_run("AI 八大维度合规审核 · 废标风险预警"), "仿宋", 10.5, color="86909C")
    doc.add_paragraph()

    # 项目信息
    def info_line(label, value):
        para = doc.add_paragraph()
        font(para.add_run(f"{label}："), "黑体", 11, True)
        font(para.add_run(str(value or "—")), "仿宋", 11)

    info_line("项目名称", pi.get("project_name"))
    info_line("采购人", pi.get("purchaser"))
    info_line("招标文件", meta.get("tender_file"))
    info_line("投标文件", meta.get("bid_file"))
    info_line("检测时间", datetime.now().strftime("%Y-%m-%d %H:%M"))

    # 分隔
    doc.add_paragraph()

    # 总览
    _heading(doc, font, "一、检测总览")
    para = doc.add_paragraph()
    font(para.add_run("预估得分："), "黑体", 12, True)
    score = report.get("overall_score")
    sc_color = "00B42A" if (isinstance(score, int) and score >= 85) else ("FF7D00" if (isinstance(score, int) and score >= 70) else "F53F3F")
    font(para.add_run(f"{score if score is not None else '—'} 分"), "黑体", 14, True, sc_color)

    para = doc.add_paragraph()
    font(para.add_run("整体结论："), "黑体", 12, True)
    font(para.add_run(report.get("overall_verdict") or "检测完成"), "仿宋", 12)
    if report.get("score_note"):
        para = doc.add_paragraph(); font(para.add_run(report["score_note"]), "仿宋", 11, color="86909C")

    para = doc.add_paragraph()
    font(para.add_run(f"否决项 {counts.get('red',0)} 项"), "仿宋", 12, True, "F53F3F")
    font(para.add_run("    "), "仿宋", 12)
    font(para.add_run(f"扣分项 {counts.get('yellow',0)} 项"), "仿宋", 12, True, "FF7D00")
    font(para.add_run("    "), "仿宋", 12)
    font(para.add_run(f"待核验 {counts.get('blue',0)} 项"), "仿宋", 12, True, "722ED1")
    font(para.add_run("    "), "仿宋", 12)
    font(para.add_run(f"合规 {counts.get('green',0)} 项"), "仿宋", 12, True, "00B42A")

    para = doc.add_paragraph()
    font(para.add_run(f"招标需求 {meta.get('requirement_count','—')} 条 · 投标分 {meta.get('section_count','—')} 章 · 检测用时 {meta.get('elapsed_seconds','—')}s"),
         "仿宋", 9, color="A0A0A0")

    # 问题清单
    findings = report.get("findings") or []
    _heading(doc, font, f"二、问题清单（共 {len(findings)} 条）")
    for i, f in enumerate(findings, 1):
        label, color = _sev(f.get("severity"))
        para = doc.add_paragraph()
        font(para.add_run(f"{i}. "), "黑体", 12, True)
        font(para.add_run(f"[{label}]"), "黑体", 11, True, color)
        if f.get("is_knockout"):
            font(para.add_run("[废标点]"), "黑体", 11, True, "F53F3F")
        font(para.add_run(f"[{f.get('dimension','')}] "), "仿宋", 11, color="86909C")
        font(para.add_run(f.get("title", "")), "黑体", 12, True)

        if f.get("location"):
            para = doc.add_paragraph(); font(para.add_run(f"位置：{f['location']}"), "仿宋", 10.5, color="86909C")
        if f.get("description"):
            para = doc.add_paragraph(); font(para.add_run(f["description"]), "仿宋", 11)
        if f.get("rule_reference"):
            para = doc.add_paragraph(); font(para.add_run(f"招标条款：{f['rule_reference']}"), "楷体", 10.5, color="4E5969")
        if f.get("bid_reference"):
            para = doc.add_paragraph(); font(para.add_run(f"投标原文：{f['bid_reference']}"), "楷体", 10.5, color="4E5969")
        if f.get("law_reference"):
            para = doc.add_paragraph(); font(para.add_run(f"法律依据：{f['law_reference']}"), "楷体", 10.5, color="4E5969")
        if f.get("suggestion"):
            para = doc.add_paragraph()
            font(para.add_run("修改建议："), "黑体", 11, True, "1677FF")
            font(para.add_run(f["suggestion"]), "仿宋", 11)
        for a in (f.get("actions") or []):
            para = doc.add_paragraph(style="List Bullet"); font(para.add_run(str(a)), "仿宋", 10.5)
        if f.get("score_impact"):
            para = doc.add_paragraph(); font(para.add_run(f"评分影响：{f['score_impact']}"), "仿宋", 10.5, color="86909C")
        doc.add_paragraph()

    # 修改优先级
    fp = report.get("fix_priority") or []
    if fp:
        _heading(doc, font, "三、优先修改顺序")
        for p in fp:
            para = doc.add_paragraph()
            font(para.add_run(f"{p.get('rank','')}. "), "黑体", 11, True, "1677FF")
            font(para.add_run(f"{p.get('id','')}  "), "黑体", 11, True)
            font(para.add_run(p.get("reason", "")), "仿宋", 11)

    # 免责
    doc.add_paragraph()
    para = doc.add_paragraph(); para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    font(para.add_run("本报告由 AI 生成，仅供参考，最终以人工复核为准。"), "仿宋", 9, color="A0A0A0")

    buf = io.BytesIO(); doc.save(buf); return buf.getvalue()


def _heading(doc, font, text):
    from docx.shared import Pt
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(10)
    font(para.add_run(text), "黑体", 15, True, "1D2129")


# ════════════════════════════════════════════════════════
#  PDF (reportlab)
# ════════════════════════════════════════════════════════
_FONT_REGULAR = "CNSong"
_FONT_BOLD = "CNHei"
_fonts_ready = False


def _first_existing(candidates):
    """candidates: [(path, ttc_subfont_index)]，返回第一个存在的。"""
    for path, idx in candidates:
        if path and os.path.exists(path):
            return path, idx
    return None, 0


def _register(name, path, idx):
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    if path.lower().endswith(".ttc"):
        pdfmetrics.registerFont(TTFont(name, path, subfontIndex=idx or 0))
    else:
        pdfmetrics.registerFont(TTFont(name, path))


def _ensure_fonts():
    """跨平台注册中文字体：本地 Windows 用系统字体；Linux/Render 用容器内 Noto/WQY 或随包字体。"""
    global _fonts_ready
    if _fonts_ready:
        return

    win = os.path.join(os.environ.get("WINDIR", "C:\\Windows"), "Fonts")
    bundled = os.path.join(os.path.dirname(__file__), "..", "fonts")

    regular = [
        (os.path.join(bundled, "chinese.ttf"), 0),
        (os.path.join(bundled, "chinese.ttc"), 0),
        ("/usr/share/fonts/truetype/wqy/wqy-microhei.ttc", 0),
        ("/usr/share/fonts/truetype/wqy/wqy-zenhei.ttc", 0),
        ("/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc", 0),
        ("/usr/share/fonts/opentype/noto/NotoSansCJKsc-Regular.otf", 0),
        (os.path.join(win, "simsun.ttc"), 0),
        (os.path.join(win, "msyh.ttc"), 0),
    ]
    bold = [
        (os.path.join(bundled, "chinese-bold.ttf"), 0),
        ("/usr/share/fonts/truetype/wqy/wqy-zenhei.ttc", 0),
        ("/usr/share/fonts/truetype/wqy/wqy-microhei.ttc", 0),
        ("/usr/share/fonts/opentype/noto/NotoSansCJK-Bold.ttc", 0),
        (os.path.join(win, "simhei.ttf"), None),
        (os.path.join(win, "simsun.ttc"), 0),
    ]

    rp, ridx = _first_existing(regular)
    if not rp:
        raise RuntimeError("未找到可用中文字体（PDF 导出需要）")
    bp, bidx = _first_existing(bold)
    if not bp:
        bp, bidx = rp, ridx

    _register(_FONT_REGULAR, rp, ridx)
    _register(_FONT_BOLD, bp, bidx)
    _fonts_ready = True


def build_pdf(payload: dict) -> bytes:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import mm
    from reportlab.lib import colors
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, HRFlowable, Table, TableStyle,
    )

    _ensure_fonts()
    pi, report, meta, counts = _overview_lines(payload)

    def esc(s):
        return (str(s) if s is not None else "").replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")

    body = ParagraphStyle("body", fontName=_FONT_REGULAR, fontSize=10.5, leading=16, textColor=colors.HexColor("#1D2129"))
    small = ParagraphStyle("small", parent=body, fontSize=9, textColor=colors.HexColor("#86909C"))
    h1 = ParagraphStyle("h1", fontName=_FONT_BOLD, fontSize=15, leading=22, spaceBefore=12, spaceAfter=6, textColor=colors.HexColor("#1D2129"))
    title = ParagraphStyle("title", fontName=_FONT_BOLD, fontSize=22, leading=30, alignment=1, textColor=colors.HexColor("#1D2129"))
    subtitle = ParagraphStyle("subtitle", fontName=_FONT_REGULAR, fontSize=10, alignment=1, textColor=colors.HexColor("#86909C"))

    flow = []
    flow.append(Paragraph("标书规范性检测报告", title))
    flow.append(Paragraph("AI 八大维度合规审核 · 废标风险预警", subtitle))
    flow.append(Spacer(1, 8 * mm))

    # 项目信息表
    info_rows = [
        ["项目名称", pi.get("project_name") or "—", "采购人", pi.get("purchaser") or "—"],
        ["招标文件", meta.get("tender_file") or "—", "投标文件", meta.get("bid_file") or "—"],
        ["检测时间", datetime.now().strftime("%Y-%m-%d %H:%M"),
         "数据量", f"需求{meta.get('requirement_count','—')}条/分{meta.get('section_count','—')}章"],
    ]
    info_rows = [[Paragraph(esc(c), small if i % 2 else ParagraphStyle('k', parent=small, fontName=_FONT_BOLD, textColor=colors.HexColor("#4E5969"))) for i, c in enumerate(r)] for r in info_rows]
    tbl = Table(info_rows, colWidths=[24 * mm, 60 * mm, 24 * mm, 52 * mm])
    tbl.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#F7F8FA")),
        ("BOX", (0, 0), (-1, -1), 0.5, colors.HexColor("#E8ECF0")),
        ("INNERGRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#E8ECF0")),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("TOPPADDING", (0, 0), (-1, -1), 5), ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
        ("LEFTPADDING", (0, 0), (-1, -1), 6), ("RIGHTPADDING", (0, 0), (-1, -1), 6),
    ]))
    flow.append(tbl)
    flow.append(Spacer(1, 6 * mm))

    # 总览
    flow.append(Paragraph("一、检测总览", h1))
    score = report.get("overall_score")
    sc_color = "#00B42A" if (isinstance(score, int) and score >= 85) else ("#FF7D00" if (isinstance(score, int) and score >= 70) else "#F53F3F")
    flow.append(Paragraph(
        f'预估得分：<font name="{_FONT_BOLD}" size="15" color="{sc_color}">{score if score is not None else "—"} 分</font>', body))
    flow.append(Paragraph(f"整体结论：{esc(report.get('overall_verdict') or '检测完成')}", body))
    if report.get("score_note"):
        flow.append(Paragraph(esc(report["score_note"]), small))
    flow.append(Paragraph(
        f'<font color="#F53F3F">● 否决项 {counts.get("red",0)} 项</font>　'
        f'<font color="#FF7D00">● 扣分项 {counts.get("yellow",0)} 项</font>　'
        f'<font color="#722ED1">● 待核验 {counts.get("blue",0)} 项</font>　'
        f'<font color="#00B42A">● 合规 {counts.get("green",0)} 项</font>', body))
    flow.append(Spacer(1, 3 * mm))
    flow.append(HRFlowable(width="100%", thickness=0.6, color=colors.HexColor("#E8ECF0")))

    # 问题清单
    findings = report.get("findings") or []
    flow.append(Paragraph(f"二、问题清单（共 {len(findings)} 条）", h1))
    for i, f in enumerate(findings, 1):
        label, color = _sev(f.get("severity"))
        ko = f' <font name="{_FONT_BOLD}" color="#F53F3F">[废标点]</font>' if f.get("is_knockout") else ""
        flow.append(Paragraph(
            f'<font name="{_FONT_BOLD}">{i}.</font> '
            f'<font name="{_FONT_BOLD}" color="#{color}">[{label}]</font>{ko} '
            f'<font color="#86909C">[{esc(f.get("dimension"))}]</font> '
            f'<font name="{_FONT_BOLD}">{esc(f.get("title"))}</font>', body))
        if f.get("location"):
            flow.append(Paragraph(f"位置：{esc(f['location'])}", small))
        if f.get("description"):
            flow.append(Paragraph(esc(f["description"]), body))
        if f.get("rule_reference"):
            flow.append(Paragraph(f'<font color="#4E5969">招标条款：{esc(f["rule_reference"])}</font>', small))
        if f.get("bid_reference"):
            flow.append(Paragraph(f'<font color="#4E5969">投标原文：{esc(f["bid_reference"])}</font>', small))
        if f.get("law_reference"):
            flow.append(Paragraph(f'<font color="#4E5969">法律依据：{esc(f["law_reference"])}</font>', small))
        if f.get("suggestion"):
            flow.append(Paragraph(f'<font name="{_FONT_BOLD}" color="#1677FF">修改建议：</font>{esc(f["suggestion"])}', body))
        for a in (f.get("actions") or []):
            flow.append(Paragraph(f"· {esc(a)}", small))
        if f.get("score_impact"):
            flow.append(Paragraph(f'<font color="#86909C">评分影响：{esc(f["score_impact"])}</font>', small))
        flow.append(Spacer(1, 4 * mm))

    # 优先修改
    fp = report.get("fix_priority") or []
    if fp:
        flow.append(Paragraph("三、优先修改顺序", h1))
        for p in fp:
            flow.append(Paragraph(
                f'<font name="{_FONT_BOLD}" color="#1677FF">{esc(p.get("rank"))}.</font> '
                f'<font name="{_FONT_BOLD}">{esc(p.get("id"))}</font>　{esc(p.get("reason"))}', body))

    flow.append(Spacer(1, 8 * mm))
    flow.append(Paragraph("本报告由 AI 生成，仅供参考，最终以人工复核为准。", subtitle))

    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4, leftMargin=18 * mm, rightMargin=18 * mm, topMargin=16 * mm, bottomMargin=16 * mm)
    doc.build(flow)
    return buf.getvalue()
