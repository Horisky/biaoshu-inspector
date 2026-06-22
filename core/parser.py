"""
文档解析：PDF / Word → 纯文本（保留段落结构与标题层级）。
（移植自主平台 backend，零内部依赖，可独立使用。）
"""
import io
from pathlib import Path


def parse_pdf(data: bytes) -> str:
    """pdfplumber 提取文本，保留页码分隔符。"""
    import pdfplumber
    pages = []
    with pdfplumber.open(io.BytesIO(data)) as pdf:
        for i, page in enumerate(pdf.pages, start=1):
            text = page.extract_text(x_tolerance=2, y_tolerance=2) or ""
            if text.strip():
                pages.append(f"--- 第{i}页 ---\n{text.strip()}")
    return "\n\n".join(pages)


def parse_docx(data: bytes) -> str:
    """python-docx 提取段落文本，标题段落加 # 前缀辅助章节识别。"""
    from docx import Document
    doc = Document(io.BytesIO(data))
    lines = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        if para.style.name.startswith("Heading"):
            level = para.style.name.replace("Heading ", "").strip()
            lines.append(f"{'#' * int(level) if level.isdigit() else '#'} {text}")
        else:
            lines.append(text)

    for table in doc.tables:
        table_lines = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            table_lines.append(" | ".join(cells))
        if table_lines:
            lines.append("\n".join(table_lines))

    return "\n\n".join(lines)


def parse_document(data: bytes, filename: str) -> str:
    ext = Path(filename).suffix.lower()
    if ext == ".pdf":
        return parse_pdf(data)
    elif ext in (".docx", ".doc"):
        return parse_docx(data)
    else:
        raise ValueError(f"不支持的文件格式: {ext}（支持 .pdf / .docx）")


def split_into_sections(bid_text: str, max_chars_per_section: int = 4000) -> list[dict]:
    """
    把投标文件正文切成章节列表，供审核引擎逐章对照。
    优先按 Markdown 标题（# 开头，来自 docx Heading）切分；
    没有标题时按长度兜底分块。返回 [{order_index, title, content}]。
    """
    lines = bid_text.split("\n")
    sections: list[dict] = []
    cur_title = "投标文件"
    cur_buf: list[str] = []

    def flush():
        content = "\n".join(cur_buf).strip()
        if content or len(sections) == 0:
            sections.append({
                "order_index": len(sections) + 1,
                "title": cur_title,
                "content": content,
            })

    has_heading = any(ln.lstrip().startswith("#") for ln in lines)
    if has_heading:
        for ln in lines:
            stripped = ln.lstrip()
            if stripped.startswith("#"):
                if cur_buf:
                    flush()
                    cur_buf = []
                cur_title = stripped.lstrip("#").strip() or "未命名章节"
            else:
                cur_buf.append(ln)
        flush()
    else:
        # 无标题：按长度分块
        chunk: list[str] = []
        size = 0
        idx = 1
        for para in bid_text.split("\n\n"):
            if size + len(para) > max_chars_per_section and chunk:
                sections.append({
                    "order_index": idx,
                    "title": f"投标文件 第{idx}部分",
                    "content": "\n\n".join(chunk).strip(),
                })
                idx += 1
                chunk = [para]
                size = len(para)
            else:
                chunk.append(para)
                size += len(para)
        if chunk:
            sections.append({
                "order_index": idx,
                "title": f"投标文件 第{idx}部分" if idx > 1 else "投标文件",
                "content": "\n\n".join(chunk).strip(),
            })

    # 过滤纯空章节，但至少保留一个
    non_empty = [s for s in sections if s["content"]]
    return non_empty or sections
